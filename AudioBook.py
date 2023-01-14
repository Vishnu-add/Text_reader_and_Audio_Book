import pyttsx3
import PyPDF2
from tkinter.filedialog import *
import docx2txt
import docx
#from gtts import gTTS

speaker = pyttsx3.init()
#   CHANGING  VOICE
voices = speaker.getProperty('voices')
speaker.setProperty('voice', voices[0].id)
#    CHANGING RATE
rate = speaker.getProperty('rate')
speaker.setProperty('rate', rate)
count=0



# def in_lelo():
#     speaker.say("To change the voice,Press zero or 1 or 2");print("To change the voice,Press 0 or 1 or 2")
#     speaker.say("I'm zero,1 is Alexa,2 is Natasha");print("I'm zero,1 is Alexa,2 is Natasha")
#
    # return vo




def shuru_karo():
    # in_lelo()
    # speaker.setProperty('voice', voices[vo].id)
    speaker.say("Select the file:");print("Select the file:")
    speaker.runAndWait()
    read_file(askopenfilename())

def read_file(path):
    #   For file in pdf format
    if path.endswith(".pdf"):
        with open(path,'rb') as book:
            global count
            print("This is in .pdf format")
            speaker.say("This is in .pdf format")
            speaker.runAndWait()
            print(" Press 1 to read out,Press 2 to save to mp3: ")
            speaker.say("Press 1 to read out,Press 2 to save to mp3: ")
            speaker.runAndWait()
            n = int(input())
            pdfReader = PyPDF2.PdfFileReader(book)
            pagesno = pdfReader.numPages
            #-----------For Entering the correct option and for the first time---------------------
            if count==0:
                #   Selecting to read out the pdf file
                if n == 1:
                    print("Total number of pages are: ", pagesno)
                    speaker.say("Total number of pages are :" + str(pagesno))
                    speaker.runAndWait()
                    print("Enter the page number you want to read out or enter '0' to read the whole pdf and E to exit:")
                    speaker.say("Enter the page number you want to read out or enter '0' to read the whole pdf and E to exit:")
                    speaker.runAndWait()
                    pno=input()
                    #    To get into the program
                    if pno.isdigit():
                        pno=int(pno)
                        #    Reading the Whole pdf file
                        if pno == 0:
                            print("Reading the whole PDF,Listen Idiot")
                            speaker.say("Reading the whole PDF,Listen Idiot")
                            speaker.runAndWait()
                            #   Reading each page for each iteration
                            for num in range(0, pagesno):
                                page = pdfReader.getPage(num)
                                text = page.extractText()
                                speaker.say(text)
                                speaker.runAndWait()
                                speaker.stop()
                            print("Done Reading")
                            speaker.say("Done Reading")
                            speaker.runAndWait()
                        #    Reading a Specific page in  a pdf file
                        elif pno != 0 and pno <= pagesno:
                            print("Listen Idiot,Reading page number:",pno)
                            speaker.say("Listen Idiot,Reading page number:"+str(pno))
                            speaker.runAndWait()
                            #for num in range(0, pagesno):
                            page = pdfReader.getPage(pno-1)
                            text = page.extractText()
                            speaker.say(text)
                            speaker.runAndWait()
                            speaker.stop()
                            print("Done Reading")
                            speaker.say("Done Reading")
                            speaker.runAndWait()
                        #    Selecting a wrong page number
                        elif pno < 0 or pno > pagesno:
                            print("Check the number of pages Idiot!!!")
                            speaker.say("Check the number of pages Idiot!!!")
                            speaker.runAndWait()
                            print("Enter the correct page number")
                            speaker.say("Enter the correct page number")
                            speaker.runAndWait()
                            count+=1
                            read_file(path, n)
                    #    For Exiting
                    elif pno == 'e' or 'E':
                        print("Exiting")
                        speaker.say("Exiting")
                        speaker.runAndWait()
                        exit()
                #   Selecting to convert the file to mp3
                elif n == 2:
                    print("Total number of pages are: ", pagesno)
                    speaker.say("Total number of pages are :" + str(pagesno))
                    speaker.runAndWait()
                    print("Enter the page number you want to Convert into or enter '0' to Convert the whole pdf and E to exit::")
                    speaker.say("Enter the page number you want to Convert into or enter '0' to Convert the whole pdf and E to exit::")
                    speaker.runAndWait()
                    ppno = input()
                    #    To get into the program
                    if ppno.isdigit():
                        ppno=int(ppno)
                        #    For coverting the whole pdf file
                        if ppno==0:
                            print("Converting to mp3")
                            speaker.say("Converting to mp3")
                            speaker.runAndWait()
                            for num in range(0, pagesno):
                                page = pdfReader.getPage(num)
                                text = page.extractText()
                                # speaker.say(text)
                                # gtts_transformer = gTTS(text=text, lang='en')
                                # gtts_transformer.save("Testttt.mp3")
                                speaker.save_to_file(text, 'test2.mp3')
                                speaker.runAndWait()
                                speaker.stop()
                            print("Done")
                            speaker.say("Done")
                            speaker.runAndWait()
                        #    For converting a specified page number
                        elif ppno!=0 and ppno<pagesno:
                            print("Converting page number ",ppno," to mp3")
                            speaker.say("Converting page number "+str(ppno)+" to mp3")
                            speaker.runAndWait()
                            #for num in range(0, pagesno):
                            page = pdfReader.getPage(ppno)
                            text = page.extractText()
                            # speaker.say(text)
                            # gtts_transformer = gTTS(text=text, lang='en')
                            # gtts_transformer.save("Testttt.mp3")
                            aname = input("Enter a name for your Audio File: ")
                            speaker.say("Enter a name for your Audio File: ")
                            speaker.runAndWait()
                            speaker.save_to_file(text, aname+'.mp3')
                            speaker.runAndWait()
                            speaker.stop()

                            print("Done")
                            speaker.say("Done")
                            speaker.runAndWait()
                        #    Selecting a wrong page number
                        elif ppno<0 or ppno>pagesno:
                            print("Check the number of pages Idiot!!!")
                            speaker.say("Check the number of pages Idiot!!!")
                            speaker.runAndWait()
                            print("Enter the correct page number")
                            speaker.say("Enter the correct page number")
                            speaker.runAndWait()
                            count += 1
                            read_file(path, n)
                    #     For Exiting
                    elif ppno=='e' or 'E':
                        print("Exiting")
                        speaker.say("Exiting")
                        speaker.runAndWait()
                        exit()
                #   Exiting
                elif n==0:
                    exit()
                #   Selecting the wrong option rather than readout and converting
                else:
                    print("Invalid option,we will add your option ASSSSSSSHOOOOOOLLLE!!!")
                    speaker.say("Invalid Option,we will add your option ASSHOLE!!!")
                    speaker.runAndWait()

            #-----------For entering the wrong option and and visiting 'n' times calling again the function to read th correct page no---------------------------------------------
            elif count != 0:
                #   Selecting to read out the pdf file
                if n == 1:
                    print("Enter the page number you want to read out or enter '0' to read the whole pdf and E to exit:")
                    speaker.say(
                        "Enter the page number you want to read out or enter '0' to read the whole pdf and E to exit:")
                    speaker.runAndWait()
                    pno = input()
                    #    To get into the program
                    if pno.isdigit():
                        pno = int(pno)
                        #    Reading the Whole pdf file
                        if pno == 0:
                            print("Reading the whole PDF,Listen Idiot")
                            speaker.say("Reading the whole PDF,Listen Idiot")
                            speaker.runAndWait()
                            #   Reading each page for each iteration
                            for num in range(0, pagesno):
                                page = pdfReader.getPage(num)
                                text = page.extractText()
                                speaker.say(text)
                                speaker.runAndWait()
                                speaker.stop()
                            print("Done Reading")
                            speaker.say("Done Reading")
                            speaker.runAndWait()
                        #    Reading a Specific pafe in  a pdf file
                        elif pno != 0 and pno <= pagesno:
                            print("Listen Idiot,Reading page number:", pno)
                            speaker.say("Listen Idiot,Reading page number:" + str(pno))
                            speaker.runAndWait()
                            # for num in range(0, pagesno):
                            page = pdfReader.getPage(pno - 1)
                            text = page.extractText()
                            speaker.say(text)
                            speaker.runAndWait()
                            speaker.stop()
                            print("Done Reading")
                            speaker.say("Done Reading")
                            speaker.runAndWait()
                        #    Selecting a wrong page number
                        elif pno < 0 or pno > pagesno:
                            print("Check the number of pages Idiot!!!")
                            speaker.say("Check the number of pages Idiot!!!")
                            speaker.runAndWait()
                            print("Enter the correct page number")
                            speaker.say("Enter the correct page number")
                            speaker.runAndWait()
                            count += 1
                            read_file(path, n)
                    #    For Exiting
                    elif pno == 'e' or 'E':
                        print("Exiting")
                        speaker.say("Exiting")
                        speaker.runAndWait()
                        exit()
                #   Selecting to convert the file to mp3
                elif n == 2:
                    print("Enter the page number you want to Convert into or enter '0' to Convert the whole pdf and E to exit::")
                    speaker.say("Enter the page number you want to Convert into or enter '0' to Convert the whole pdf and E to exit::")
                    speaker.runAndWait()
                    ppno = input()
                    #    To get into the program
                    if ppno.isdigit():
                        ppno = int(ppno)
                        #    For coverting the whole pdf file
                        if ppno == 0:
                            print("Converting to mp3")
                            speaker.say("Converting to mp3")
                            speaker.runAndWait()
                            for num in range(0, pagesno):
                                page = pdfReader.getPage(num)
                                text = page.extractText()
                                # speaker.say(text)
                                # gtts_transformer = gTTS(text=text, lang='en')
                                # gtts_transformer.save("Testttt.mp3")
                                speaker.save_to_file(text, 'test2.mp3')
                                speaker.runAndWait()
                                speaker.stop()
                            print("Done")
                            speaker.say("Done")
                            speaker.runAndWait()
                        #    For converting a specified page number
                        elif ppno != 0 and ppno < pagesno:
                            print("Converting page number ", ppno, " to mp3")
                            speaker.say("Converting page number " + str(ppno) + " to mp3")
                            speaker.runAndWait()
                            # for num in range(0, pagesno):
                            page = pdfReader.getPage(ppno)
                            text = page.extractText()
                            # speaker.say(text)
                            # gtts_transformer = gTTS(text=text, lang='en')
                            # gtts_transformer.save("Testttt.mp3")
                            aname = input("Enter a name for your Audio File: ")
                            speaker.say("Enter a name for your Audio File: ")
                            speaker.runAndWait()
                            speaker.save_to_file(text, aname + '.mp3')
                            speaker.runAndWait()
                            speaker.stop()

                            print("Done")
                            speaker.say("Done")
                            speaker.runAndWait()
                        #    Selecting a wrong page number
                        elif ppno < 0 or ppno > pagesno:
                            print("Check the number of pages Idiot!!!")
                            speaker.say("Check the number of pages Idiot!!!")
                            speaker.runAndWait()
                            print("Enter the correct page number")
                            speaker.say("Enter the correct page number")
                            speaker.runAndWait()
                            count += 1
                            read_file(path, n)
                    #     For Exiting
                    elif ppno == 'e' or 'E':
                        print("Exiting")
                        speaker.say("Exiting")
                        speaker.runAndWait()
                        exit()
                #   Selecting the wrong option rather than readout and converting
                else:
                    print("Invalid option,we will add your option ASSSSSSSHOOOOOOLLLE!!!")
                    speaker.say("Invalid Option,we will add your option ASSHOLE!!!")
                    speaker.runAndWait()
        path.close()
    #   For file in txt format
    elif path.endswith(".txt"):
        with open(path,'r') as file:
            print("This is in .txt format")
            speaker.say("This is in .txt format")
            speaker.runAndWait()
            print("Press 1 to read out,Press 2 to save to mp3,Press 0 to exit: ")
            speaker.say("Press 1 to read out,Press 2 to save to mp3,Press 0 to exit: ")
            speaker.runAndWait()
            n = int(input())
            #   Reading the text file
            if n==1:
                line=file.read()
                speaker.say(line)
                speaker.runAndWait()
                print("Done Reading")
                speaker.say("Done Reading")
                speaker.runAndWait()
            #   Selecting to convert the file to mp3
            elif n==2:
                line=file.read()
                speaker.save_to_file(line,'text.mp3')
                speaker.runAndWait()
                print("Done Converting")
                speaker.say("Done Converting")
                speaker.runAndWait()
            #   Exiting
            elif n == 0:
                print("Exiting")
                speaker.say("Exiting")
                speaker.runAndWait()
                exit()
            #   Selecting the wrong option rather than readout and converting
            else:
                print("Invalid option,we will add your option ASSSSSSSHOOOOOOLLLE!!!")
                speaker.say("Invalid Option,we will add your option ASSHOLE!!!")
                speaker.runAndWait()
        path.close()
    #   For file in docx format
    elif path.endswith('.docx'):
        print("This is in .docx format")
        speaker.say("This is in .docx format")
        speaker.runAndWait()
        print("Press 1 to read out,Press 2 to save to mp3,Press 0 to exit: ")
        speaker.say("Press 1 to read out,Press 2 to save to mp3,Press 0 to exit: ")
        speaker.runAndWait()
        n = int(input())
        #   Reading the word doc file
        if n == 1:
            print("Listen Idiot")
            speaker.say("Listen Idiot")
            speaker.runAndWait()
            # doc = docx.Document(path)
            # all_paras = doc.paragraphs
            # wordtext = []
            # for para in all_paras:
            #     wordtext.append(para.text)
            wordtext = docx2txt.process(path)
            speaker.say(wordtext)
            speaker.runAndWait()
            print("Done Reading")
            speaker.say("Done Reading")
            speaker.runAndWait()
        #    Converting to mp3
        elif n == 2:
            print("Converting")
            speaker.say("Converting")
            speaker.runAndWait()
            # doc = docx.Document(path)
            # all_paras = doc.paragraphs
            # wordtext=[]
            # for para in all_paras:
            #     wordtext.append(para.text)
            wordtext = docx2txt.process(path)
            speaker.save_to_file(wordtext,"word.mp3")
            speaker.runAndWait()
            print("Done Converting")
            speaker.say("Done Converting")
            speaker.runAndWait()
        #    Exiting
        elif n == 0:
            print("Exiting")
            speaker.say("Exiting")
            speaker.runAndWait()
            exit()
    #   For file other than the specified formats
    else:
        print("This is'nt a pdf file")


# in_lelo()
# vo = int(input())
shuru_karo()
